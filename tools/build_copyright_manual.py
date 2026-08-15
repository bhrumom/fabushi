from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "copyright_registration" / "发布软件V1.0_操作说明书.docx"
SHOTS = ROOT / "output" / "copyright_registration" / "screenshots"
BLUE = "2E74B5"
LIGHT = "E8EEF5"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.addnext(fld)
    paragraph.add_run(" 页")


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)


def add_step(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)


def add_screenshot(doc: Document, filename: str, caption: str) -> None:
    path = SHOTS / filename
    if not path.exists():
        raise FileNotFoundError(f"Missing screenshot: {path}")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(6.35))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.style = doc.styles["Caption"]


def new_page(doc: Document, title: str) -> None:
    doc.add_page_break()
    doc.add_heading(title, level=1)


def build() -> None:
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, before, after, color in (
        ("Title", 30, 0, 18, BLUE),
        ("Heading 1", 16, 18, 10, BLUE),
        ("Heading 2", 13, 14, 7, BLUE),
        ("Heading 3", 12, 10, 5, "1F4D78"),
    ):
        st = doc.styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)

    for name in ("List Bullet", "List Number"):
        st = doc.styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
        st.paragraph_format.left_indent = Inches(0.375)
        st.paragraph_format.first_line_indent = Inches(-0.188)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.line_spacing = 1.25

    cap = doc.styles["Caption"]
    cap.font.name = "Calibri"
    cap.font.size = Pt(9)
    cap.font.color.rgb = RGBColor.from_string("5B6573")
    cap._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")

    header = sec.header.paragraphs[0]
    header.text = "发布软件 V1.0  ·  操作说明书"
    header.style = doc.styles["Caption"]
    footer = sec.footer.paragraphs[0]
    add_page_number(footer)
    footer.style = doc.styles["Caption"]

    # Cover
    for _ in range(4):
        doc.add_paragraph()
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("发布软件 V1.0")
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("操 作 说 明 书")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor.from_string("1F4D78")
    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    desc.add_run("AI 对话、小程序宿主、运行控制与通信平台").italic = True
    for _ in range(7):
        doc.add_paragraph()
    owner = doc.add_paragraph()
    owner.alignment = WD_ALIGN_PARAGRAPH.CENTER
    owner.add_run("著作权人：广西谛曦人工智能应用软件有限公司")
    date = doc.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date.add_run("文档版本：V1.0    编制日期：2026年8月13日")

    new_page(doc, "1  文档信息与软件概述")
    doc.add_heading("1.1 文档信息", level=2)
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(1.75)
    table.columns[1].width = Inches(4.75)
    set_repeat_table_header(table.rows[0])
    for i, text in enumerate(("项目", "内容")):
        cell = table.rows[0].cells[i]
        cell.text = text
        set_cell_shading(cell, LIGHT)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    rows = [
        ("软件名称", "发布软件"),
        ("版本号", "V1.0"),
        ("软件类型", "macOS桌面应用软件"),
        ("主要用途", "AI对话、小程序安装运行、权限审批、任务控制及会话通信"),
        ("开发完成日期", "2026年8月13日"),
        ("运行平台", "macOS 12及以上"),
    ]
    for a, b in rows:
        cells = table.add_row().cells
        cells[0].text, cells[1].text = a, b
    for row in table.rows:
        for cell in row.cells:
            set_cell_margins(cell)

    doc.add_heading("1.2 软件定位", level=2)
    doc.add_paragraph(
        "发布软件是一套采用Tauri 2、Rust和React/TypeScript构建的macOS桌面端应用平台。"
        "软件以AI对话为统一入口，提供Marketplace小程序安装与打开、隔离运行、能力授权、"
        "长任务中断、会话清理和运行状态反馈，为AI小程序的创作、交付与通信流程提供安全宿主。"
    )
    doc.add_heading("1.3 功能组成", level=2)
    for text in (
        "AI对话：发送自然语言任务并接收Rust运行时返回的流式事件与结果。",
        "小程序宿主：从Marketplace安装并打开小程序，在隔离容器中运行。",
        "交付衔接：通过小程序清单、安装状态和打开状态衔接生成与发布流程。",
        "通信协作：管理会话消息及运行事件，保留连续任务上下文。",
        "安全运行：能力授权确认、长任务中断、会话清理、异常反馈和状态记录。",
    ):
        add_bullet(doc, text)

    new_page(doc, "2  启动软件与主界面")
    doc.add_heading("2.1 启动条件", level=2)
    add_bullet(doc, "建议使用Apple Silicon或Intel Mac，8GB内存及以上，并保留足够的本地存储空间。")
    add_bullet(doc, "需要在线模型、同步或发布功能时，应保持网络连接；本地功能可按实际配置离线运行。")
    doc.add_heading("2.2 启动步骤", level=2)
    add_step(doc, "在macOS中打开发布软件应用。")
    add_step(doc, "等待初始化完成，进入桌面主界面。")
    add_step(doc, "根据工作目标使用AI对话、小程序、权限审批或运行控制区域。")
    add_screenshot(doc, "01-home.png", "图2-1  发布软件 macOS桌面端主界面")

    new_page(doc, "3  AI对话与任务指令")
    doc.add_heading("3.1 发送任务", level=2)
    add_step(doc, "在主界面的聊天区域输入任务内容。")
    add_step(doc, "描述小程序用途、界面结构、交互逻辑或待处理事项。")
    add_step(doc, "点击“发送”，查看用户消息及运行时返回的结果。")
    doc.add_paragraph(
        "AI对话、消息事件和运行结果均在主界面的会话区域展示；相关区域已包含在图2-1的最新版实机主界面截图中。"
    )
    doc.add_heading("3.2 任务反馈", level=2)
    doc.add_paragraph(
        "系统通过版本化命令协议把任务交给Rust运行时，并将消息、增量、完成或失败事件返回界面。"
        "用户可以在同一会话中继续补充要求；当小程序调用受控能力时，软件显示审批界面并等待用户决策。"
    )
    doc.add_heading("3.3 结果检查", level=2)
    for text in (
        "确认消息已显示，并检查宿主状态为ready。",
        "检查运行结果、错误提示及对应操作编号。",
        "发现错误时保留原会话并重试；必要时清理会话后重新建立任务。",
    ):
        add_bullet(doc, text)
    add_screenshot(doc, "05-running-task.png", "图3-1  AI任务运行与状态控制界面")

    new_page(doc, "4  Marketplace小程序安装与运行")
    doc.add_paragraph(
        "Marketplace、MiniApp入口和安装状态位于主界面的小程序区域；相关入口已包含在图2-1的最新版实机主界面截图中。"
    )
    doc.add_heading("4.1 安装与打开", level=2)
    add_step(doc, "在Marketplace与MiniApp区域选择目标小程序。")
    add_step(doc, "点击“安装”，等待状态由not-installed变为installed。")
    add_step(doc, "点击“打开”，进入隔离MiniApp容器。")
    add_screenshot(doc, "02-marketplace.png", "图4-1  Marketplace插件市场与扩展分类")
    add_screenshot(doc, "03-miniapp.png", "图4-2  MiniApp安装后在宿主中的运行界面")
    doc.add_heading("4.2 预览与运行", level=2)
    doc.add_paragraph(
        "小程序打开后，宿主显示其标识、容器状态和可申请的能力。用户可通过状态输出检查安装、打开和运行链路。"
        "若运行失败，应先查看错误信息，再检查小程序标识、运行时状态、权限和网络条件。"
    )

    new_page(doc, "5  能力审批与安全控制")
    doc.add_paragraph(
        "能力申请由小程序在受控宿主内触发，系统在审批卡中显示能力名称、调用来源和申请原因；"
        "只有用户作出允许或拒绝决定后，运行时才继续执行对应受控调用。"
    )
    doc.add_heading("5.1 发起能力申请", level=2)
    add_step(doc, "在已打开的小程序容器中点击能力申请按钮。")
    add_step(doc, "查看申请能力、小程序标识和使用原因。")
    add_step(doc, "仅在确认任务可信且确有必要时选择允许一次，否则选择拒绝。")
    doc.add_heading("5.2 审批结果", level=2)
    add_step(doc, "宿主把审批决策传递给Rust运行时。")
    add_step(doc, "查看审批状态是否变为allowed或denied。")
    add_step(doc, "审批仅用于当前受控调用，不代表对所有小程序永久授权。")
    doc.add_paragraph(
        "注意：涉及系统权限、外部平台登录、验证码、授权或最终发布确认时，应由用户依据对应规则亲自完成。"
    )

    new_page(doc, "6  运行控制与会话管理")
    doc.add_heading("6.1 长任务控制", level=2)
    add_step(doc, "点击“启动长任务”，观察操作状态变为running。")
    add_step(doc, "需要停止时点击“中断”，等待运行时返回interrupted事件。")
    add_step(doc, "查看操作编号和最终状态，确认任务已安全结束。")
    doc.add_heading("6.2 会话管理", level=2)
    add_bullet(doc, "聊天消息和运行事件构成当前会话上下文。")
    add_bullet(doc, "点击“清理会话”可清除当前运行时会话状态并重新开始。")
    add_bullet(doc, "账号与联系人协议为后续跨端通信提供统一数据结构。")
    add_screenshot(doc, "05-computer.png", "图6-1  Computer与运行状态面板")
    add_screenshot(doc, "06-automation.png", "图6-2  自动化例程创建与管理界面")

    new_page(doc, "7  权限、安全与运行记录")
    doc.add_heading("7.1 权限控制", level=2)
    doc.add_paragraph(
        "小程序和插件通过受控宿主调用系统能力。涉及文件、网络、外部应用或其他敏感能力时，软件依据权限配置限制调用范围。"
        "用户应仅为可信任务授予必要权限，并在任务结束后检查授权状态。"
    )
    doc.add_heading("7.2 日志与故障恢复", level=2)
    for text in (
        "任务状态用于显示消息、安装、打开、审批、运行和会话过程。",
        "异常记录包含阶段、错误信息和必要的上下文，便于回溯问题。",
        "可恢复任务支持重试或重新进入；重要内容建议在发布前保留本地副本。",
    ):
        add_bullet(doc, text)
    doc.add_heading("7.3 数据使用建议", level=2)
    add_bullet(doc, "不要在普通对话或草稿中输入不必要的账号密码、密钥或个人敏感信息。")
    add_bullet(doc, "上传文件或向外部平台发布前，应确认内容权属和目标范围。")
    add_bullet(doc, "使用第三方模型、插件或发布渠道时，应同时遵守对应服务的规则。")
    add_screenshot(doc, "07-settings.png", "图7-1  设置与本机能力管理界面")

    new_page(doc, "8  常见问题与退出")
    doc.add_heading("8.1 常见问题", level=2)
    qa = [
        ("小程序无法加载", "检查小程序是否已启用、入口配置是否有效，并查看运行状态或错误日志。"),
        ("消息任务中断", "确认宿主状态和运行时连接，保留原会话后重试；必要时清理会话。"),
        ("能力申请失败", "确认小程序已打开、能力名称有效，并根据审批状态重新操作。"),
        ("文件无法读取", "确认文件仍在原位置、格式受支持，并在macOS系统设置中检查必要的文件访问权限。"),
        ("运行状态不同步", "检查宿主连接，重新进入会话或执行会话清理。"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(1.9)
    table.columns[1].width = Inches(4.6)
    for i, text in enumerate(("问题", "处理方法")):
        table.rows[0].cells[i].text = text
        set_cell_shading(table.rows[0].cells[i], LIGHT)
    set_repeat_table_header(table.rows[0])
    for q, a in qa:
        cells = table.add_row().cells
        cells[0].text, cells[1].text = q, a
    for row in table.rows:
        for cell in row.cells:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_heading("8.2 正常退出", level=2)
    add_step(doc, "确认正在执行的消息或长任务已完成或已安全中断。")
    add_step(doc, "确认小程序和审批状态已记录。")
    add_step(doc, "关闭发布软件窗口；如需切换账号，应先按界面提示退出当前账号。")
    doc.add_paragraph("—— 文档结束 ——").alignment = WD_ALIGN_PARAGRAPH.CENTER

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
